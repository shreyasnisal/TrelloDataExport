#### HOW TO USE THIS SCRIPT ####

### This script generates a single word document with data from multiple trello boards

## Follow the below steps to generate the word document

# Export the trello boards in JSON format
# Put all the JSON files in the same directory as this script
# Run this script using the command `python create_trello_doc.py`

import urllib.request
from PIL import Image
import json
import os
from docx import Document
from docx.shared import Inches
import docx

@property
def image_width(self):
    if (self.horz_dpi == 0):
        return Inches(self.px_width / 72)
    return Inches(self.px_width / self.horz_dpi)


@property
def image_height(self):
    if (self.vert_dpi == 0):
        return Inches(self.px_height / 72)
    return Inches(self.px_height / self.vert_dpi)


docx.image.image.Image.width = image_width
docx.image.image.Image.height = image_height

DIR_PATH = "."

for filename in os.listdir(DIR_PATH):
   
   if not filename.endswith(".json"):
      continue

   word_doc = Document()

   print("Converting file {}".format(filename))

   word_doc.add_heading(filename, level=1)

   with open(os.path.join(DIR_PATH, filename), 'r', encoding='utf-8') as json_file:
      json_data = json.load(json_file)
      lists = {}
      for list in json_data["lists"]:
         lists[list["id"]] = list["name"]

      list_cards = {}
      for card in json_data['cards']:
         card_data = {"name": card["name"], "desc": card["desc"], "attachments": []}

         for attachment in card["attachments"]:
            if attachment['url'].endswith('png') or attachment['url'].endswith('jpg') or attachment['url'].endswith('jpeg'):
               card_data["attachments"].append(attachment["url"])
         
         if lists[card["idList"]] in list_cards:
            list_cards[lists[card["idList"]]].append(card_data)
         else:
            list_cards[lists[card["idList"]]] = [card_data]

      for list in list_cards:
         word_doc.add_heading(list, level=2)
         for card in list_cards[list]:
            word_doc.add_heading(card["name"], level=3)
            word_doc.add_paragraph(card["desc"])
            for attachment in card["attachments"]:
               image_name = attachment.split('/')[-1]
               urllib.request.urlretrieve(attachment, image_name)
               word_doc.add_picture(image_name, width=Inches(4.0))

   word_doc.save(filename.split('.')[0] + '.docx')
